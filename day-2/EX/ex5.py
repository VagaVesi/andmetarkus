# salvestada tulemus wordi

import pandas as pd
from docx import Document
import requests

# Assuming 'df' is your DataFrame with the results
# If not, replace 'df' with your actual DataFrame variable

url_rahvastik_ = 'https://api.worldbank.org/v2/country/EST/indicator/SP.POP.TOTL?format=json'
url_women = "https://api.worldbank.org/v2/country/EST/indicator/SP.POP.TOTL.FE.IN?format=json"

response = requests.get(url_rahvastik_)
data = response.json()

response_woman = requests.get(url_women)
data_woman = response_woman.json()

values = {'year': [], 'population': [], 'women_population': []}


for item in data[1]:
    values['year'].append(item['date'])
    values['population'].append(item['value'])

for women in data_woman[1]:
    values['women_population'].append(women['value'])

# print(json.dumps(values, indent=2, ensure_ascii=False))

df = pd.DataFrame(values)

df = df.sort_values(by='year')  # sorteerin aasta järgi kasvavalt

# Create a new Word document
doc = Document()
doc.add_heading('Eesti rahvaarv aastatel 1960-2021', 0)

# Add DataFrame as a table
table = doc.add_table(rows=1, cols=len(df.columns))
hdr_cells = table.rows[0].cells
for i, col_name in enumerate(df.columns):
    hdr_cells[i].text = col_name

for index, row in df.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)

# Save the document
doc.save("rahvaarv_tulemus.docx")
